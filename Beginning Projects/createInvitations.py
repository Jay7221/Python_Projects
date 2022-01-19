#! python3
# createInvitations.py - Makes word document invitations for a given guest list.

guestList=['Alice','Bob','Tom','Donny']
fonts=['Custom','Custom','Custom','Custom','Custom']
import docx
doc=docx.Document('Invitations.docx')
for guestName in guestList:
    doc.add_paragraph('It would be a pleasure to have the company of'.center(100),fonts[0])
    doc.add_paragraph(guestName.center(100),fonts[1])
    doc.add_paragraph('at 11010 Memory Lane on the Evening of'.center(100),fonts[2])
    doc.add_paragraph('April 1st'.center(100),fonts[3])
    doc.add_paragraph('at 7 o\'clock'.center(100),fonts[4])
    #doc.paragraphs[4].runs[0].add_break(doc.text.WD_BREAK.PAGE)

doc.save('Invitations.docx')
print('Done!')
